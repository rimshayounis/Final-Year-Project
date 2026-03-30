from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ─── Styles helpers ──────────────────────────────────────────────────────────
def heading(text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
    return p

def para(text, bold=False, italic=False, size=None, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    return p

def table_row(tbl, cells, bold_first=False):
    row = tbl.add_row()
    for i, (cell, val) in enumerate(zip(row.cells, cells)):
        cell.text = str(val)
        if bold_first and i == 0:
            for run in cell.paragraphs[0].runs:
                run.bold = True
    return row

def add_table(headers, rows, bold_first_col=False):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = 'Table Grid'
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
    for row in rows:
        table_row(tbl, row, bold_first=bold_first_col)
    doc.add_paragraph()
    return tbl

# ─── Title Page ──────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("MedConnect — Financial & Feature Model")
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x4B, 0x00, 0x82)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Version 4 — March 2026")
r2.font.size = Pt(12)
r2.italic = True

doc.add_paragraph()
para("This document describes MedConnect's monetisation strategy, reward ecosystem, "
     "mentor progression framework, and personalised recommendation engine as implemented "
     "in the current codebase.", size=11)
doc.add_page_break()

# ─── Section 1: Overview ─────────────────────────────────────────────────────
heading("1. Platform Overview", level=1)
para("MedConnect is a telehealth platform connecting patients with verified doctors through "
     "appointments, a social health feed, and AI-assisted care tools.")
bullet("Patients: book appointments, follow health feeds, access SOS features")
bullet("Doctors: earn revenue via appointments, build reputation through the Mentor Level system")
bullet("Admins: verify doctors, manage subscriptions, monitor analytics")
doc.add_paragraph()

# ─── Section 2: Revenue Streams ──────────────────────────────────────────────
heading("2. Revenue Streams", level=1)

heading("2.1 Subscription Plans (Doctors)", level=2)
add_table(
    ["Plan", "Price / Month", "Verification Slots", "Features"],
    [
        ["Free Trial", "PKR 0", "1", "Basic profile, limited appointments"],
        ["Basic", "PKR 999", "3", "Standard analytics, appointment management"],
        ["Premium", "PKR 2,499", "10", "Priority listing, advanced analytics, badge"],
        ["Enterprise", "PKR 5,999", "Unlimited", "White-label profile, API access, dedicated support"],
    ],
    bold_first_col=True,
)

heading("2.2 Appointment Commissions", level=2)
bullet("Platform retains 10–15 % of each appointment fee paid through the in-app wallet")
bullet("Payments held in escrow; released automatically when appointment session ends")
bullet("Auto-release handled by NestJS cron scheduler (checks every minute)")
doc.add_paragraph()

heading("2.3 Premium Features (Patients)", level=2)
bullet("Priority appointment booking (skip queue)")
bullet("Unlimited chat history export")
bullet("Personalised AI health reports (planned)")
doc.add_paragraph()

# ─── Section 3: Trust Score (Internal) ───────────────────────────────────────
heading("3. Trust Score — Internal Reputation Signal", level=1)
para("Trust Score is an internal integer (0–4) awarded to doctors based on cumulative post "
     "engagement milestones. It is NOT displayed to users; it feeds the Mentor Level formula.")

add_table(
    ["Milestone (total likes)", "Trust Score awarded"],
    [
        ["First post reaches 10 likes", "+1"],
        ["Cumulative 50 likes", "+1"],
        ["Cumulative 200 likes", "+1"],
        ["Cumulative 500 likes", "+1"],
    ],
)

para("Trust Score is stored in the PointsReward wallet document. Every time a post receives "
     "a like, the scheduler re-evaluates milestones and updates the score atomically.", italic=True)
doc.add_paragraph()

# ─── Section 4: Points & Rewards ─────────────────────────────────────────────
heading("4. Points & Rewards System", level=1)
para("Doctors earn redeemable points for posting quality health content.")

add_table(
    ["Action", "Points Earned", "Notes"],
    [
        ["Post approved by admin", "10 pts", "Awarded once per post"],
        ["Post receives 10 likes", "25 pts", "Milestone bonus"],
        ["Post receives 50 likes", "50 pts", "Milestone bonus"],
        ["Post receives 200 likes", "100 pts", "Milestone bonus"],
        ["Appointment completed", "5 pts", "Per confirmed completion"],
    ],
)

heading("4.1 Redemption", level=2)
bullet("Points convert to wallet credit at PKR 0.50 per point")
bullet("Minimum redemption threshold: 100 points")
bullet("Redeemed balance offsets next subscription renewal")
doc.add_paragraph()

# ─── Section 5: Mentor Level ─────────────────────────────────────────────────
heading("5. Mentor Level System", level=1)
para("Mentor Level is a visible, composite progression metric shown on doctor profiles "
     "and appointment cards. It replaces the retired Trust Badge display.")

heading("5.1 Formula", level=2)
para("MentorScore = (TrustScore x 25) + CompletedAppointments + round(AvgRating x 10)", bold=True)
doc.add_paragraph()
bullet("TrustScore — internal signal (0–4) from like milestones (Section 3)")
bullet("CompletedAppointments — total sessions marked COMPLETED in the database")
bullet("AvgRating — mean star rating across all reviewed appointments (1.0–5.0)")
doc.add_paragraph()

heading("5.2 Level Thresholds", level=2)
add_table(
    ["Level", "Title", "Min Score", "Display Colour"],
    [
        ["Lv. 1", "Newcomer", "0", "Grey"],
        ["Lv. 2", "Rising", "25", "Green"],
        ["Lv. 3", "Trusted", "60", "Blue"],
        ["Lv. 4", "Expert", "120", "Purple"],
        ["Lv. 5", "Master", "230", "Gold / Amber"],
    ],
    bold_first_col=True,
)

heading("5.3 Example Calculation", level=2)
para("Dr. Aisha: TrustScore=2, CompletedAppointments=18, AvgRating=4.3")
para("  MentorScore = (2 x 25) + 18 + round(4.3 x 10) = 50 + 18 + 43 = 111  →  Lv. 3 Trusted", bold=True)
doc.add_paragraph()

heading("5.4 Display Locations", level=2)
bullet("Doctor Profile Screen — ribbon icon + level title + score progress bar")
bullet("Book Appointment card — coloured chip showing Lv.N Title")
bullet("Admin Doctor Detail page — level summary panel")
doc.add_paragraph()

# ─── Section 6: Interest-Based Recommendation ────────────────────────────────
heading("6. Interest-Based Recommendation System", level=1)
para("Patients select health interests during onboarding (Health Profile screen). "
     "These interests drive personalisation across the feed and the appointment booking flow.")

heading("6.1 Interest Categories", level=2)
add_table(
    ["Interest Label", "Mapped Post Categories / Specializations"],
    [
        ["Skin Care", "Dermatology, Beauty"],
        ["Hair Care", "Dermatology, Trichology"],
        ["Weight Loss", "Nutrition, Fitness, Diet"],
        ["Weight Gain", "Nutrition, Fitness, Diet"],
        ["Mental Health", "Psychiatry, Psychology, Mental Wellness"],
        ["Heart Health", "Cardiology, Cardiovascular"],
        ["Diabetes", "Endocrinology, Internal Medicine"],
        ["General Fitness", "Fitness, Exercise, Lifestyle"],
    ],
    bold_first_col=True,
)

heading("6.2 Feed Recommendation Algorithm", level=2)
para("Endpoint: GET /posts/feed/recommended?userId=<id>&page=N")
bullet("Step 1 — Load patient interests from healthProfile.interests")
bullet("Step 2 — Map interests to post category keywords")
bullet("Step 3 — Find other users sharing at least one interest (same-interest peers)")
bullet("Step 4 — MongoDB aggregation scoring:")
bullet("+2 points if post category matches a mapped category", level=1)
bullet("+1 point if post author is a same-interest peer", level=1)
bullet("Step 5 — Sort by score DESC, then createdAt DESC")
bullet("Step 6 — Paginate and return (isPersonalized flag included in response)")
para("Fallback: if patient has no interests, returns chronological feed.", italic=True)
doc.add_paragraph()

heading("6.3 Doctor Recommendation (Book Appointment)", level=2)
para("No additional backend endpoint — pure frontend sort applied in BookAppointment.tsx.")
bullet("Patient interests mapped to specialization keywords (INTEREST_SPEC_MAP)")
bullet("Each doctor's specialization string checked for keyword matches")
bullet("Matched doctors sorted to the top of the list")
bullet("Matched doctors display a 'Suggested for you' badge (purple, sparkle icon)")
bullet("Tie-breaking: subscription plan rank, then completedCount")
doc.add_paragraph()

# ─── Section 7: Wallet & Escrow ───────────────────────────────────────────────
heading("7. Wallet & Escrow Flow", level=1)
add_table(
    ["Step", "Actor", "Action"],
    [
        ["1", "Patient", "Pays appointment fee into escrow wallet"],
        ["2", "System", "Holds funds; appointment session begins"],
        ["3", "Scheduler", "Every minute: checks if session end time passed"],
        ["4", "Scheduler", "Marks appointment COMPLETED; triggers releaseAppointmentPayment"],
        ["5", "System", "Transfers net amount to doctor wallet; platform retains commission"],
        ["6", "Doctor", "Can request withdrawal or apply balance to subscription"],
    ],
)
doc.add_paragraph()

# ─── Section 8: SOS & Safety Features ────────────────────────────────────────
heading("8. SOS & Patient Safety", level=1)
bullet("Shake-to-trigger SOS via accelerometer (expo-sensors)")
bullet("On trigger: captures GPS coordinates (expo-location)")
bullet("Dispatches emergency email with location link via SOS service")
bullet("Alert shown in-app with countdown and cancel option")
doc.add_paragraph()

# ─── Section 9: Admin Revenue Controls ───────────────────────────────────────
heading("9. Admin Module", level=1)
bullet("Doctor verification & subscription management")
bullet("Post moderation (approve / reject) — approval triggers point award")
bullet("Analytics dashboard: revenue, appointments, active users")
bullet("Mentor Level visible per doctor in detail panel")
bullet("Subscription plan assignment and plan-change audit log")
doc.add_paragraph()

# ─── Section 10: Financial Projections ───────────────────────────────────────
heading("10. Financial Projections (12-Month Estimate)", level=1)
add_table(
    ["Metric", "Month 3", "Month 6", "Month 12"],
    [
        ["Active Doctors", "50", "200", "600"],
        ["Avg Subscription Revenue / Doctor (PKR)", "1,200", "1,500", "1,800"],
        ["Subscription Revenue (PKR)", "60,000", "300,000", "1,080,000"],
        ["Appointment Commission Revenue (PKR)", "20,000", "150,000", "500,000"],
        ["Total Monthly Revenue (PKR)", "80,000", "450,000", "1,580,000"],
    ],
    bold_first_col=True,
)
para("Projections assume 15 % MoM doctor growth and 10 appointments/doctor/month at avg PKR 800.", italic=True)
doc.add_paragraph()

# ─── Section 11: Complete Feature Summary ────────────────────────────────────
heading("11. Complete Feature Summary", level=1)
add_table(
    ["Feature", "Status", "Notes"],
    [
        ["Doctor Registration & Verification", "Done", "Multi-step form, admin approval"],
        ["Patient Registration & Health Profile", "Done", "Includes interests selection"],
        ["Appointment Booking", "Done", "Slot-based, calendar UI"],
        ["Real-time Chat (in-appointment)", "Done", "Session-gated WebSocket"],
        ["Auto Chat-Close & Payment Release", "Done", "Cron scheduler, every minute"],
        ["Post Feed", "Done", "Category tabs + interest-ranked All tab"],
        ["Interest-Based Post Recommendation", "Done", "Score aggregation pipeline"],
        ["Interest-Based Doctor Recommendation", "Done", "Frontend sort + badge"],
        ["SOS Emergency Feature", "Done", "Shake trigger, GPS, email dispatch"],
        ["Points & Rewards", "Done", "Milestones, redemption logic"],
        ["Trust Score (Internal)", "Done", "Hidden from UI; feeds Mentor Level"],
        ["Mentor Level System", "Done", "Formula-based, 5 levels, visible on profile"],
        ["Subscription Plans", "Done", "4 tiers, admin-managed"],
        ["Wallet & Escrow", "Done", "Escrow hold, auto-release"],
        ["Admin Analytics Dashboard", "Done", "Revenue, users, appointments"],
        ["AI Health Reports", "Planned", "Future roadmap item"],
    ],
    bold_first_col=True,
)

# ─── Footer ───────────────────────────────────────────────────────────────────
doc.add_page_break()
p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_footer = p_footer.add_run("MedConnect — Confidential — Version 4 — March 2026")
r_footer.italic = True
r_footer.font.size = Pt(9)
r_footer.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.save("financial_model_v4.docx")
print("financial_model_v4.docx generated successfully.")
